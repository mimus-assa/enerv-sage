from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from formato_contratos import formato_distribuidor
from roman import convertir_a_romano

class ConfiguradorDocumento:
    def __init__(self, documento):
        self.documento = documento

    def configurar_pagina(self):
        sections = self.documento.sections
        for section in sections:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.left_margin = Inches(1.18)
            section.right_margin = Inches(1.18)
            section.top_margin = Inches(0.49)
            section.bottom_margin = Inches(0.49)

    def configurar_fuente(self, parrafo, fuente="Arial", tamano_fuente=12):
        for run in parrafo.runs:
            run.font.name = fuente
            run.font.size = Pt(tamano_fuente)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), fuente)
        parrafo.style.font.name = fuente
        parrafo.style.font.size = Pt(tamano_fuente)

class ProcesadorTexto:
    def __init__(self, documento):
        self.documento = documento

    def agregar_texto_con_formato(self, parrafo, texto, tamano_fuente, color=None, negrita=False):
        run = parrafo.add_run(texto)
        run.font.size = Pt(tamano_fuente)
        run.font.name = "Arial"
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), "Arial")
        if negrita:
            run.bold = True
        if color:
            run.font.color.rgb = RGBColor(color[0], color[1], color[2])

    def procesar_etiquetas(self, texto, tamano_fuente):
        partes = []
        i = 0
        while i < len(texto):
            if texto[i] == '<':
                end_tag_pos = texto.find('>', i) + 1
                tag = texto[i:end_tag_pos]
                
                if tag.startswith('<b>'):
                    end_bold_tag_pos = texto.find('</b>', i) + len('</b>')
                    bold_text = texto[i + len('<b>'):end_bold_tag_pos - len('</b>')]
                    partes.append(('text', bold_text, tamano_fuente, None, True))
                    i = end_bold_tag_pos
                
                elif tag.startswith('<color='):
                    end_color_tag_pos = texto.find('</color>', i) + len('</color>')
                    color_code = tag[len('<color='):tag.find('>')]
                    color_value = tuple(map(int, color_code.split(',')))
                    color_text_start = end_tag_pos
                    color_text_end = texto.find('</color>', color_text_start)
                    color_text = texto[color_text_start:color_text_end]

                    if '<b>' in color_text:
                        start_bold_in_color = color_text.find('<b>') + len('<b>')
                        end_bold_in_color = color_text.find('</b>')
                        partes.append(('text', color_text[:start_bold_in_color - len('<b>')], tamano_fuente, color_value, False))
                        partes.append(('text', color_text[start_bold_in_color:end_bold_in_color], tamano_fuente, color_value, True))
                        partes.append(('text', color_text[end_bold_in_color + len('</b>'):], tamano_fuente, color_value, False))
                    else:
                        partes.append(('text', color_text, tamano_fuente, color_value, False))
                    
                    i = end_color_tag_pos
                
                else:
                    i = end_tag_pos
            else:
                next_tag_pos = texto.find('<', i)
                if next_tag_pos == -1:
                    partes.append(('text', texto[i:], tamano_fuente, None, False))
                    break
                else:
                    partes.append(('text', texto[i:next_tag_pos], tamano_fuente, None, False))
                    i = next_tag_pos
        return partes

class DocumentoContrato:
    def __init__(self):
        self.documento = Document()
        self.documento._body.clear_content()
        self.configurador = ConfiguradorDocumento(self.documento)
        self.procesador = ProcesadorTexto(self.documento)
        self.configurador.configurar_pagina()
        self.estilo_personalizado = self.crear_estilo_personalizado('EstiloPersonalizado', 9)  # Ajusta el tamaño de fuente aquí

    def crear_estilo_personalizado(self, nombre_estilo, tamano_fuente):
        estilos = self.documento.styles
        estilo_parrafo = estilos.add_style(nombre_estilo, WD_STYLE_TYPE.PARAGRAPH)
        estilo_fuente = estilo_parrafo.font
        estilo_fuente.size = Pt(tamano_fuente)
        estilo_fuente.name = 'Arial'
        return nombre_estilo

    def agregar_parrafo(self, texto, config):
        p = self.documento.add_paragraph()
        p_format = p.paragraph_format

        if config['indentacion'] > 0:
            p_format.left_indent = Pt(config['indentacion'])

        p_format.space_before = Pt(0)
        p_format.space_after = Pt(config.get('espaciado_despues', 0))
        p_format.line_spacing = Pt(12)

        self.configurador.configurar_fuente(p, tamano_fuente=config['tamano_fuente'])

        lineas = texto.split('\n')
        for i, linea in enumerate(lineas):
            if i > 0:
                p.add_run().add_break()
            partes = self.procesador.procesar_etiquetas(linea, config['tamano_fuente'])
            for tipo, texto, tamano_fuente, color, negrita in partes:
                if tipo == 'text':
                    self.procesador.agregar_texto_con_formato(p, texto, tamano_fuente, color, negrita)

        if config.get('justificado', False):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            if config['alineacion'] == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif config['alineacion'] == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def agregar_parrafo_en_blanco(self, config):
        p = self.documento.add_paragraph()
        p.style = self.documento.styles[self.estilo_personalizado]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(config.get('espaciado_despues', 0))

    def agregar_lista(self, items, config):
        for i, item in enumerate(items):
            if item.strip():
                n = i + 1
                numero_romano = convertir_a_romano(n)
                p = self.documento.add_paragraph()

                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.5)
                p_format.first_line_indent = Inches(-0.5)
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(config.get('espaciado_despues', 0))
                p_format.line_spacing = Pt(12)

                self.configurador.configurar_fuente(p, tamano_fuente=config['tamano_fuente'])

                run = p.add_run(f"{numero_romano}. \t")
                run.font.size = Pt(config['tamano_fuente'])
                run.font.name = "Arial"

                partes = self.procesador.procesar_etiquetas(item, config['tamano_fuente'])
                for tipo, texto, tamano_fuente, color, negrita in partes:
                    if tipo == 'text':
                        self.procesador.agregar_texto_con_formato(p, texto, tamano_fuente, color, negrita)

                if config.get('justificado', False):
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    if config['alineacion'] == 'center':
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif config['alineacion'] == 'right':
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                self.agregar_parrafo_en_blanco(config)

    def procesar_linea(self, linea, configuracion, i, lineas):
        config = configuracion.get(i, {'tamano_fuente': 12, 'alineacion': 'left', 'justificado': False, 'indentacion': 0})
        if linea.startswith('<list>'):
            items = []
            i += 1
            while i < len(lineas) and not lineas[i].startswith('</list>'):
                item = lineas[i].strip().replace('<item>', '').replace('</item>', '')
                items.append(item)
                i += 1
            self.agregar_lista(items, config)
        elif linea.startswith('</list>'):
            return i + 1
        else:
            self.agregar_parrafo(linea, config)
        return i + 1

    def generar_documento(self, lineas, variables, configuracion):
        i = 0
        while i < len(lineas):
            linea = lineas[i].rstrip()
            if linea:
                for clave, valor in variables.items():
                    linea = linea.replace(f"{{{clave}}}", valor)
                i = self.procesar_linea(linea, configuracion, i, lineas)
            else:
                config = configuracion.get(i, {'tamano_fuente': 12})
                self.agregar_parrafo_en_blanco(config)
                i += 1

    def guardar_documento(self, nombre_archivo):
        self.documento.save(nombre_archivo)

def cargar_lineas_archivo(ruta_archivo):
    with open(ruta_archivo, "r", encoding="utf-8") as file:
        return file.readlines()

# Ejemplo de uso
if __name__ == "__main__":
    ruta_lineas = "mnt/data/contrato_distribuidor.txt"
    lineas = cargar_lineas_archivo(ruta_lineas)

    variables = {
        "nombre_cliente": "John Doe",
        "nombre_gestor": "Jane Smith",
        "No_INE": "1234567890123",
        "promedio_mensual": "5000",
        "voltaje_tension": "220V",
        "RPU": "987654321",
        "RMU": "567890123",
        "cuenta": "1234567890",
        "direccion_cliente": "123 Main Street, City, Country",
        "tarifa": "1C",
        "numero_hilos": "3",
        "numero_medidor": "5432109876",
        "demanda_contratada": "50 kW",
        "demanda_instalada": "60 kW",
        "telefono_cliente": "555-1234",
        "correo_cliente": "cliente@example.com",
        "lugar_fecha": "Ciudad, Fecha",
        "No_INE_gestor": "0987654321098"
    }

    doc_contrato = DocumentoContrato()
    doc_contrato.generar_documento(lineas, variables, formato_distribuidor)
    doc_contrato.guardar_documento("mnt/data/contrato_generado.docx")
